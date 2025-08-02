import json
from docx import Document
from docx.shared import RGBColor
import re

def load_json(filepath):
    with open(filepath, 'r') as file:
        return json.load(file)

def flatten_json(y, prefix=''):
    out = {}
    for key, val in y.items():
        new_key = f"{prefix}.{key}" if prefix else key
        if isinstance(val, dict):
            out.update(flatten_json(val, new_key))
        else:
            out[new_key] = val
            out[key] = val
    return out

def is_red(run):
    color = run.font.color
    return color and (color.rgb == RGBColor(255, 0, 0) or getattr(color, "theme_color", None) == 1)

def get_value_as_string(value, field_name=""):
    if isinstance(value, list):
        if len(value) == 0:
            return ""
        elif len(value) == 1:
            return str(value[0])
        else:
            if "australian company number" in field_name.lower() or "company number" in field_name.lower():
                return value
            else:
                return " ".join(str(v) for v in value)
    else:
        return str(value)

def find_matching_json_value(field_name, flat_json):
    """Find matching JSON value based on field name (key)"""
    field_name = field_name.strip()
    
    # Manual mapping for specific sections that need special handling
    manual_mappings = {
        "attendance list name and position title": "Attendance List (Names and Position Titles).Attendance List (Names and Position Titles)",
        "attendance list (names and position titles)": "Attendance List (Names and Position Titles).Attendance List (Names and Position Titles)",
        "nature of the operators business (summary)": "Nature of the Operators Business (Summary).Nature of the Operators Business (Summary)",
        "nature of the operators business (summary):": "Nature of the Operators Business (Summary).Nature of the Operators Business (Summary)",
        "nature of operators business (summary)": "Nature of the Operators Business (Summary).Nature of the Operators Business (Summary)",
        "nature of operators business (summary):": "Nature of the Operators Business (Summary).Nature of the Operators Business (Summary)",
        # Paragraph-level mappings
        "mass management": "paragraphs.MASS MANAGEMENT",
        "liam herbig": "paragraphs.MASS MANAGEMENT",  # Name should be replaced with company name
        "date": "paragraphs.This management system I have audited when followed will ensure compliance with the relevant NHVAS Business Rules & Standards.",
        # Date-related mappings
        "13.11.2024": "paragraphs.This management system I have audited when followed will ensure compliance with the relevant NHVAS Business Rules & Standards.",
        "auditor signature": "paragraphs.This management system I have audited when followed will ensure compliance with the relevant NHVAS Business Rules & Standards.",
        "operator signature": "paragraphs.I hereby consent to information relating to my Accreditation to be shared with other law enforcement agencies, including a service provider authorised under the Heavy Vehicle National Law.",
        # Specific data mappings
        "jodie jones": "Audit Information.Auditor name",
        "13th november 2024": "Audit Information.Date of Audit",
        "adelaide barossa transport & warehousing pty ltd": "Operator Information.Operator name (Legal entity)",
        "manager": "Operator Information.Operator name (Legal entity)",  # Replace manager title with company name
        "liam herbig –manager": "Operator Information.Operator name (Legal entity)",
        "liam herbig – manager": "Operator Information.Operator name (Legal entity)",
        "deborah herbig – manager": "Operator Information.Operator name (Legal entity)",
        # Contact information mappings (old data in red text -> new data from JSON)
        "141 sitz road callington sa 5254": "Operator Information.Operator business address",  # Replace old address with new
        "po box 743 mt barker sa": "Operator Information.Operator Postal address",  # Replace old postal with new
        "debherbig@bigpond.com": "Operator Information.Email address",  # Replace old email with new
        "0447 710 602": "Operator Information.Operator Telephone Number",  # Replace old phone with new
        # Manual/Version mappings (old version -> new version)
        "mahlo 092021v1": "Operator Information.NHVAS Manual (Policies and Procedures) developed by",  # Replace old manual with new
        # These should stay as they are (no replacement needed, just different format)
        "511840": "Operator Information.NHVAS Accreditation No. (If applicable)",  # Keep accreditation number
        "26th october 2023": "Audit Information.Date of Audit",  # Use audit date instead
        # Std 5 and Std 6 mappings
        "the latest verification was dated 23rdnovember 2022": "Mass Management Summary of Audit findings.Std 5. Verification",
        "the latest verification was dated 23rd november 2022": "Mass Management Summary of Audit findings.Std 5. Verification",
        "internal review was dated 23rd august 2023 with 0 ncr": "Mass Management Summary of Audit findings.Std 6. Internal Review",
        "23rd august2023 with 0 trips, 0 trips using mass, 0 overloads and 0 ncr's": "Mass Management Summary of Audit findings.Std 6. Internal Review",
        "23rd august 2023 with 0 trips, 0 trips using mass, 0 overloads and 0 ncr's": "Mass Management Summary of Audit findings.Std 6. Internal Review",
    }
    
    # Check manual mappings first
    normalized_field = field_name.lower().strip()
    if normalized_field in manual_mappings:
        mapped_key = manual_mappings[normalized_field]
        if mapped_key in flat_json:
            print(f"    ✅ Manual mapping found for '{field_name}' -> '{mapped_key}'")
            return flat_json[mapped_key]
    
    # Try exact match first
    if field_name in flat_json:
        print(f"    Direct match found for key '{field_name}'")
        return flat_json[field_name]
    
    # Try case-insensitive exact match
    for key, value in flat_json.items():
        if key.lower() == field_name.lower():
            print(f"    Case-insensitive match found for key '{field_name}' with JSON key '{key}'")
            return value
    
    # Try to find a key that ends with this field name
    for key, value in flat_json.items():
        if key.endswith('.' + field_name):
            print(f"    Suffix match found for key '{field_name}' with JSON key '{key}'")
            return value
    
    # Try partial matching for fields with parentheses or additional text
    clean_field = re.sub(r'\s*\([^)]*\)', '', field_name).strip()  # Remove parentheses content
    for key, value in flat_json.items():
        clean_key = re.sub(r'\s*\([^)]*\)', '', key).strip()
        if clean_field.lower() == clean_key.lower():
            print(f"    Clean match found for key '{field_name}' with JSON key '{key}'")
            return value
    
    # Try word-based matching - more flexible approach
    field_words = set(word.lower() for word in re.findall(r'\b\w+\b', field_name) if len(word) > 2)
    best_match = None
    best_score = 0
    
    for key, value in flat_json.items():
        key_words = set(word.lower() for word in re.findall(r'\b\w+\b', key) if len(word) > 2)
        # Calculate how many words match
        common_words = field_words.intersection(key_words)
        if common_words:
            score = len(common_words) / max(len(field_words), len(key_words))  # Normalized score
            if score > best_score:
                best_score = score
                best_match = (key, value)
    
    if best_match and best_score >= 0.5:  # At least 50% word overlap
        print(f"    Word-based match found for key '{field_name}' with JSON key '{best_match[0]}' (score: {best_score:.2f})")
        return best_match[1]
    
    # No match found
    print(f"    ❌ No match found for '{field_name}'")
    return None

def get_clean_text(cell):
    text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            text += run.text
    return text.strip()

def has_red_text(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if is_red(run) and run.text.strip():
                return True
    return False

def replace_red_text_in_cell(cell, replacement_text):
    replacements_made = 0
    
    # First, collect all red text to show what we're replacing
    all_red_text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if is_red(run):
                all_red_text += run.text
    
    if all_red_text.strip():
        print(f"      ✅ Replacing red text: '{all_red_text[:50]}...' → '{replacement_text[:50]}...'")
    
    # Now replace all red text in the cell with the replacement text
    first_replacement_done = False
    for paragraph in cell.paragraphs:
        red_runs = [run for run in paragraph.runs if is_red(run)]
        if red_runs:
            if not first_replacement_done:
                # Replace the first red run with our text
                red_runs[0].text = replacement_text
                red_runs[0].font.color.rgb = RGBColor(0, 0, 0)
                first_replacement_done = True
                replacements_made = 1
            else:
                # Clear the first red run since we already replaced content
                red_runs[0].text = ''
            
            # Clear all other red runs in this paragraph
            for run in red_runs[1:]:
                run.text = ''
    
    return replacements_made

def handle_australian_company_number(row, company_numbers):
    replacements_made = 0
    for i, digit in enumerate(company_numbers):
        cell_idx = i + 1
        if cell_idx < len(row.cells):
            cell = row.cells[cell_idx]
            if has_red_text(cell):
                cell_replacements = replace_red_text_in_cell(cell, str(digit))
                replacements_made += cell_replacements
                print(f"      -> Placed digit '{digit}' in cell {cell_idx + 1}")
    return replacements_made

def handle_vehicle_registration_table(table, flat_json):
    """Handle the Vehicle Registration Numbers table with column-based data"""
    replacements_made = 0
    
    # Look for the vehicle registration data in the flattened JSON
    vehicle_section = None
    
    # Try to find the vehicle registration section
    for key, value in flat_json.items():
        if "vehicle registration numbers of records examined" in key.lower():
            if isinstance(value, dict):  # This should be the nested structure
                vehicle_section = value
                print(f"    ✅ Found vehicle data in key: '{key}'")
                break
    
    if not vehicle_section:
        # Try alternative approach - look for individual column keys
        potential_columns = {}
        for key, value in flat_json.items():
            if any(col_name in key.lower() for col_name in ["registration number", "sub-contractor", "weight verification", "rfs suspension"]):
                # Extract the column name from the flattened key
                if "." in key:
                    column_name = key.split(".")[-1]
                else:
                    column_name = key
                potential_columns[column_name] = value
        
        if potential_columns:
            vehicle_section = potential_columns
            print(f"    ✅ Found vehicle data from flattened keys: {list(vehicle_section.keys())}")
        else:
            print(f"    ❌ Vehicle registration data not found in JSON")
            return 0
    
    print(f"    ✅ Found vehicle registration data with {len(vehicle_section)} columns")
    
    # Find header row (usually row 0 or 1)
    header_row_idx = -1
    header_row = None
    
    for row_idx, row in enumerate(table.rows):
        row_text = "".join(get_clean_text(cell).lower() for cell in row.cells)
        if "registration" in row_text and "number" in row_text:
            header_row_idx = row_idx
            header_row = row
            break
    
    if header_row_idx == -1:
        print(f"    ❌ Could not find header row in vehicle table")
        return 0
    
    print(f"    ✅ Found header row at index {header_row_idx}")
    
    # Create mapping between column indices and JSON keys
    column_mapping = {}
    for col_idx, cell in enumerate(header_row.cells):
        header_text = get_clean_text(cell).strip()
        if not header_text or header_text.lower() == "no.":
            continue
            
        # Try to match header text with JSON keys
        best_match = None
        best_score = 0
        
        # Normalize header text for better matching
        normalized_header = header_text.lower().replace("(", " (").replace(")", ") ").strip()
        
        for json_key in vehicle_section.keys():
            normalized_json = json_key.lower().strip()
            
            # Try exact match first (after normalization)
            if normalized_header == normalized_json:
                best_match = json_key
                best_score = 1.0
                break
            
            # Try word-based matching
            header_words = set(word.lower() for word in normalized_header.split() if len(word) > 2)
            json_words = set(word.lower() for word in normalized_json.split() if len(word) > 2)
            
            if header_words and json_words:
                common_words = header_words.intersection(json_words)
                score = len(common_words) / max(len(header_words), len(json_words))
                
                if score > best_score and score >= 0.3:  # At least 30% match
                    best_score = score
                    best_match = json_key
            
            # Try substring matching for cases like "RegistrationNumber" vs "Registration Number"
            header_clean = normalized_header.replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
            json_clean = normalized_json.replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
            
            if header_clean in json_clean or json_clean in header_clean:
                if len(header_clean) > 5 and len(json_clean) > 5:  # Only for meaningful matches
                    substring_score = min(len(header_clean), len(json_clean)) / max(len(header_clean), len(json_clean))
                    if substring_score > best_score and substring_score >= 0.6:
                        best_score = substring_score
                        best_match = json_key
        
        if best_match:
            column_mapping[col_idx] = best_match
            print(f"      📌 Column {col_idx + 1} ('{header_text}') -> '{best_match}' (score: {best_score:.2f})")
    
    if not column_mapping:
        print(f"    ❌ No column mappings found")
        return 0
    
    # Determine how many data rows we need based on the JSON arrays
    max_data_rows = 0
    for json_key, data in vehicle_section.items():
        if isinstance(data, list):
            max_data_rows = max(max_data_rows, len(data))
    
    print(f"    📌 Need to populate {max_data_rows} data rows")
    
    # Process all required data rows
    for data_row_index in range(max_data_rows):
        table_row_idx = header_row_idx + 1 + data_row_index
        
        # Check if this table row exists, if not, add it
        if table_row_idx >= len(table.rows):
            print(f"    ⚠️ Row {table_row_idx + 1} doesn't exist - table only has {len(table.rows)} rows")
            print(f"    ➕ Adding new row for vehicle {data_row_index + 1}")
            
            # Add a new row to the table
            new_row = table.add_row()
            print(f"    ✅ Successfully added row {len(table.rows)} to the table")
            
        row = table.rows[table_row_idx]
        print(f"    📌 Processing data row {table_row_idx + 1} (vehicle {data_row_index + 1})")
        
        # Fill in data for each mapped column
        for col_idx, json_key in column_mapping.items():
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                
                # Get the data for this column and row
                column_data = vehicle_section.get(json_key, [])
                if isinstance(column_data, list) and data_row_index < len(column_data):
                    replacement_value = str(column_data[data_row_index])
                    
                    # Check if cell has red text or is empty (needs data)
                    cell_text = get_clean_text(cell)
                    if has_red_text(cell) or not cell_text.strip():
                        # If cell is empty, add the text directly
                        if not cell_text.strip():
                            cell.text = replacement_value
                            replacements_made += 1
                            print(f"      -> Added '{replacement_value}' to empty cell (column '{json_key}')")
                        else:
                            # If cell has red text, replace it
                            cell_replacements = replace_red_text_in_cell(cell, replacement_value)
                            replacements_made += cell_replacements
                            if cell_replacements > 0:
                                print(f"      -> Replaced red text with '{replacement_value}' (column '{json_key}')")
    
    return replacements_made

def handle_print_accreditation_section(table, flat_json):
    """Handle the special case of print accreditation name with 2 values"""
    replacements_made = 0
    
    # Look for the print accreditation name data
    print_data = flat_json.get("print accreditation name.print accreditation name", [])
    if not isinstance(print_data, list) or len(print_data) < 2:
        return 0
    
    name_value = print_data[0]  # "Simon Anderson"
    position_value = print_data[1]  # "Director"
    
    print(f"    📋 Print accreditation data: Name='{name_value}', Position='{position_value}'")
    
    # Find rows with "Print Name" and "Position Title"
    for row_idx, row in enumerate(table.rows):
        if len(row.cells) >= 2:
            # Check if this row has the headers
            cell1_text = get_clean_text(row.cells[0]).lower()
            cell2_text = get_clean_text(row.cells[1]).lower()
            
            if "print name" in cell1_text and "position title" in cell2_text:
                print(f"    📍 Found header row {row_idx + 1}: '{cell1_text}' | '{cell2_text}'")
                
                # Check the next row for red text to replace
                if row_idx + 1 < len(table.rows):
                    data_row = table.rows[row_idx + 1]
                    if len(data_row.cells) >= 2:
                        # Replace Print Name (first cell)
                        if has_red_text(data_row.cells[0]):
                            cell_replacements = replace_red_text_in_cell(data_row.cells[0], name_value)
                            replacements_made += cell_replacements
                            if cell_replacements > 0:
                                print(f"      ✅ Replaced Print Name: '{name_value}'")
                        
                        # Replace Position Title (second cell)  
                        if has_red_text(data_row.cells[1]):
                            cell_replacements = replace_red_text_in_cell(data_row.cells[1], position_value)
                            replacements_made += cell_replacements
                            if cell_replacements > 0:
                                print(f"      ✅ Replaced Position Title: '{position_value}'")
                
                break  # Found the section, no need to continue
    
    return replacements_made

def process_single_column_sections(cell, field_name, flat_json):
    json_value = find_matching_json_value(field_name, flat_json)
    if json_value is not None:
        replacement_text = get_value_as_string(json_value, field_name)
        if isinstance(json_value, list) and len(json_value) > 1:
            replacement_text = "\n".join(str(item) for item in json_value)
        if has_red_text(cell):
            print(f"    ✅ Replacing red text in single-column section: '{field_name}'")
            print(f"    ✅ Replacement text:\n{replacement_text}")
            cell_replacements = replace_red_text_in_cell(cell, replacement_text)
            if cell_replacements > 0:
                print(f"    -> Replaced with: '{replacement_text[:100]}...'")
                return cell_replacements
    return 0

def process_tables(document, flat_json):
    """Process tables to find key-value pairs and replace red values"""
    replacements_made = 0
    
    for table_idx, table in enumerate(document.tables):
        print(f"\n🔍 Processing table {table_idx + 1}:")
        
        # Check if this is the vehicle registration table
        table_text = ""
        for row in table.rows[:3]:  # Check first 3 rows
            for cell in row.cells:
                table_text += get_clean_text(cell).lower() + " "
        
        # Look for vehicle registration indicators (need multiple indicators to avoid false positives)
        vehicle_indicators = ["registration number", "sub-contractor", "weight verification", "rfs suspension"]
        indicator_count = sum(1 for indicator in vehicle_indicators if indicator in table_text)
        if indicator_count >= 3:  # Require at least 3 indicators to be sure it's a vehicle table
            print(f"    🚗 Detected Vehicle Registration table")
            vehicle_replacements = handle_vehicle_registration_table(table, flat_json)
            replacements_made += vehicle_replacements
            continue  # Skip normal processing for this table
        
        # Check if this is the print accreditation table
        print_accreditation_indicators = ["print name", "position title"]
        indicator_count = sum(1 for indicator in print_accreditation_indicators if indicator in table_text)
        if indicator_count >= 2:  # Require at least 2 indicators to be sure it's a print accreditation table
            print(f"    📋 Detected Print Accreditation table")
            print_accreditation_replacements = handle_print_accreditation_section(table, flat_json)
            replacements_made += print_accreditation_replacements
            continue  # Skip normal processing for this table
        
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) < 1:  # Skip empty rows
                continue
                
            # Get the key from the first column
            key_cell = row.cells[0]
            key_text = get_clean_text(key_cell)
            
            if not key_text:
                continue
            
            print(f"  📌 Row {row_idx + 1}: Key = '{key_text}'")
            
            # Check if this key exists in our JSON
            json_value = find_matching_json_value(key_text, flat_json)
            
            if json_value is not None:
                replacement_text = get_value_as_string(json_value, key_text)
                
                # Special handling for Australian Company Number
                if ("australian company number" in key_text.lower() or "company number" in key_text.lower()) and isinstance(json_value, list):
                    cell_replacements = handle_australian_company_number(row, json_value)
                    replacements_made += cell_replacements
                    
                # Handle section headers (like Attendance List, Nature of Business) where content is in next row
                elif ("attendance list" in key_text.lower() or "nature of" in key_text.lower()) and row_idx + 1 < len(table.rows):
                    print(f"    ✅ Section header detected, checking next row for content...")
                    next_row = table.rows[row_idx + 1]
                    
                    # Check all cells in the next row for red text
                    for cell_idx, cell in enumerate(next_row.cells):
                        if has_red_text(cell):
                            print(f"    ✅ Found red text in next row, cell {cell_idx + 1}")
                            # For list values, join with line breaks
                            if isinstance(json_value, list):
                                replacement_text = "\n".join(str(item) for item in json_value)
                            cell_replacements = replace_red_text_in_cell(cell, replacement_text)
                            replacements_made += cell_replacements
                            if cell_replacements > 0:
                                print(f"    -> Replaced section content with: '{replacement_text[:100]}...'")
                                
                elif len(row.cells) == 1 or (len(row.cells) > 1 and not any(has_red_text(row.cells[i]) for i in range(1, len(row.cells)))):
                    if has_red_text(key_cell):
                        cell_replacements = process_single_column_sections(key_cell, key_text, flat_json)
                        replacements_made += cell_replacements
                else:
                    for cell_idx in range(1, len(row.cells)):
                        value_cell = row.cells[cell_idx]
                        if has_red_text(value_cell):
                            print(f"    ✅ Found red text in column {cell_idx + 1}")
                            cell_replacements = replace_red_text_in_cell(value_cell, replacement_text)
                            replacements_made += cell_replacements
            else:
                if len(row.cells) == 1 and has_red_text(key_cell):
                    red_text = ""
                    for paragraph in key_cell.paragraphs:
                        for run in paragraph.runs:
                            if is_red(run):
                                red_text += run.text
                    if red_text.strip():
                        section_value = find_matching_json_value(red_text.strip(), flat_json)
                        if section_value is not None:
                            section_replacement = get_value_as_string(section_value, red_text.strip())
                            cell_replacements = replace_red_text_in_cell(key_cell, section_replacement)
                            replacements_made += cell_replacements
                
                # Handle tables where red text appears in multiple columns (like contact info tables)
                for cell_idx in range(len(row.cells)):
                    cell = row.cells[cell_idx]
                    if has_red_text(cell):
                        # Get the red text from this cell
                        red_text = ""
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if is_red(run):
                                    red_text += run.text
                        
                        if red_text.strip():
                            # Try to find a direct mapping for this red text
                            section_value = find_matching_json_value(red_text.strip(), flat_json)
                            if section_value is not None:
                                section_replacement = get_value_as_string(section_value, red_text.strip())
                                cell_replacements = replace_red_text_in_cell(cell, section_replacement)
                                replacements_made += cell_replacements
                                if cell_replacements > 0:
                                    print(f"    ✅ Replaced red text '{red_text.strip()[:30]}...' with '{section_replacement[:30]}...' in cell {cell_idx + 1}")
    
    return replacements_made

def process_paragraphs(document, flat_json):
    replacements_made = 0
    print(f"\n🔍 Processing paragraphs:")
    for para_idx, paragraph in enumerate(document.paragraphs):
        red_runs = [run for run in paragraph.runs if is_red(run) and run.text.strip()]
        if red_runs:
            full_text = paragraph.text.strip()
            red_text_only = "".join(run.text for run in red_runs).strip()
            print(f"  📌 Paragraph {para_idx + 1}: Found red text: '{red_text_only}'")
            
            # Try to match the red text specifically first
            json_value = find_matching_json_value(red_text_only, flat_json)
            
            # If no match, try some common patterns
            if json_value is None:
                # Check for signature patterns
                if "AUDITOR SIGNATURE" in red_text_only.upper() or "DATE" in red_text_only.upper():
                    json_value = find_matching_json_value("auditor signature", flat_json)
                elif "OPERATOR SIGNATURE" in red_text_only.upper():
                    json_value = find_matching_json_value("operator signature", flat_json)
                    
            if json_value is not None:
                replacement_text = get_value_as_string(json_value)
                print(f"    ✅ Replacing red text with: '{replacement_text}'")
                red_runs[0].text = replacement_text
                red_runs[0].font.color.rgb = RGBColor(0, 0, 0)
                for run in red_runs[1:]:
                    run.text = ''
                replacements_made += 1
    return replacements_made

def main():
    json_path = 'updated_word_data.json'
    docx_path = 'test.docx'
    output_path = 'updated_reportv1.docx'

    try:
        json_data = load_json(json_path)
        flat_json = flatten_json(json_data)
        print("📄 Available JSON keys (sample):")
        count = 0
        for key, value in sorted(flat_json.items()):
            if count < 10:
                print(f"  - {key}: {value}")
                count += 1
        print(f"  ... and {len(flat_json) - count} more keys\n")

        doc = Document(docx_path)

        table_replacements = process_tables(doc, flat_json)
        paragraph_replacements = process_paragraphs(doc, flat_json)
        total_replacements = table_replacements + paragraph_replacements

        doc.save(output_path)
        print(f"\n✅ Document saved as: {output_path}")
        print(f"✅ Total replacements: {total_replacements} ({table_replacements} in tables, {paragraph_replacements} in paragraphs)")

    except FileNotFoundError as e:
        print(f"❌ File not found: {e}")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
