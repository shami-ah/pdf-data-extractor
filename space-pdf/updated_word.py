#!/usr/bin/env python3
# update_docx_from_json.py
import sys, json, re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import RGBColor, Pt  # add Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

BLACK = RGBColor(0, 0, 0)
RED = RGBColor(0xFF, 0x00, 0x00)

# ----------------------------- text helpers -----------------------------
def _find_table_with_headers(doc: Document, must_have: list[str]) -> Optional[Table]:
    for t in doc.tables:
        if not t.rows: 
            continue
        head = canon(" ".join(cell_text(c) for c in t.rows[0].cells))
        if all(canon_label(x) in head for x in must_have):
            return t
    return None

def ensure_auditor_decl_headers(doc: Document) -> bool:
    """
    Second-last page table under 'NHVAS APPROVED AUDITOR DECLARATION'.
    Force the HEADER row to read exactly:
      [ Print Name | NHVR or Exemplar Global Auditor Registration Number ]
    Never touch the bottom (values) row.
    """
    changed = False
    expected_left  = "Print Name"
    expected_right = "NHVR or Exemplar Global Auditor Registration Number"

    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        # must look like the auditor table: header left says "Print Name", 2+ cols, 2+ rows
        head_left = canon_label(cell_text(t.rows[0].cells[0]))
        if head_left == "print name" and len(t.rows[0].cells) >= 2 and len(t.rows) >= 2:
            # fix left header if needed
            if canon_label(cell_text(t.rows[0].cells[0])) != canon_label(expected_left) or \
               any(is_red_run(r) for p in t.rows[0].cells[0].paragraphs for r in p.runs):
                _set_cell_text_black(t.rows[0].cells[0], expected_left)
                changed = True
            # unconditionally set the RIGHT header text (this is where "Peter Sheppard" was sitting)
            if canon_label(cell_text(t.rows[0].cells[1])) != canon_label(expected_right) or \
               any(is_red_run(r) for p in t.rows[0].cells[1].paragraphs for r in p.runs):
                _set_cell_text_black(t.rows[0].cells[1], expected_right)
                changed = True
            # found and fixed the table; no need to continue
            break

    return changed


def fill_operator_declaration(doc: Document, print_name: str, position_title: str) -> bool:
    """Last page table: write values ONLY into the bottom row (red placeholders)."""
    t = _find_table_with_headers(doc, ["Print Name", "Position Title"])
    if not t or len(t.rows) < 2 or len(t.rows[0].cells) < 2:
        return False
    bot_left  = t.rows[1].cells[0]
    bot_right = t.rows[1].cells[1]

    # only replace if that cell has a red placeholder
    if any(is_red_run(r) for p in bot_left.paragraphs for r in p.runs):
        _set_cell_text_black(bot_left, print_name)
    if any(is_red_run(r) for p in bot_right.paragraphs for r in p.runs):
        _set_cell_text_black(bot_right, position_title)
    return True

def find_heading_index_from_end(doc: Document, heading: str) -> Optional[int]:
    key = canon(heading)
    allp = iter_paragraphs(doc)
    for i in range(len(allp) - 1, -1, -1):
        if key in canon(para_text(allp[i])):
            return i
    return None

def set_date_by_heading_from_end(doc: Document, heading: str, date_text: str, max_scan: int = 60) -> bool:
    """Find the LAST occurrence of `heading`, then replace the FIRST red run in the next paragraphs."""
    if not date_text:
        return False
    allp = iter_paragraphs(doc)
    idx = find_heading_index_from_end(doc, heading)
    if idx is None:
        return False
    for p in allp[idx + 1 : min(idx + 1 + max_scan, len(allp))]:
        if replace_red_in_paragraph(p, date_text):  # writes in black
            return True
    return False

def set_date_by_paragraph_from_end(doc: Document, paragraph_text: str, date_text: str, max_scan: int = 60) -> bool:
    """Find the LAST paragraph matching `paragraph_text`, then set the FIRST red run after it."""
    if not date_text:
        return False
    key = canon(paragraph_text)
    allp = iter_paragraphs(doc)
    hit = None
    for i in range(len(allp) - 1, -1, -1):
        if key in canon(para_text(allp[i])):
            hit = i
            break
    if hit is None:
        return False
    # date placeholder is on the LAST page, right after this long paragraph
    for p in allp[hit + 1 : min(hit + 1 + max_scan, len(allp))]:
        if replace_red_in_paragraph(p, date_text):  # writes in black
            return True
    return False

def set_layer3_name_after_management_heading(doc: Document, mid_heading: str, allowed_prev_titles: List[str], name: str) -> bool:
    if not name:
        return False

    allp = iter_paragraphs(doc)
    wrote = False
    mid = canon(mid_heading)
    allowed_prev = {canon(t) for t in allowed_prev_titles}

    for i, p in enumerate(allp):
        if canon(para_text(p)) != mid:
            continue

        # previous non-empty must be one of the allowed titles
        j = i - 1
        while j >= 0 and not nz(para_text(allp[j])):
            j -= 1
        if j < 0 or canon(para_text(allp[j])) not in allowed_prev:
            continue

        # next non-empty is the 3rd line we overwrite
        k = i + 1
        while k < len(allp) and not nz(para_text(allp[k])):
            k += 1
        if k >= len(allp):
            continue

        # compute target size from the middle heading; fall back to a sensible bump
        target_size = _para_effective_font_size(allp[i]) or Pt(16)

        _clear_para_and_write_black(allp[k], name)

        # apply size to all runs explicitly (overrides style)
        for r in allp[k].runs:
            r.font.size = target_size

        wrote = True

    return wrote

def _para_effective_font_size(p: Paragraph):
    # try explicit run sizes first
    for r in p.runs:
        if r.font.size:
            return r.font.size
    # then the paragraph style
    if p.style and p.style.font and p.style.font.size:
        return p.style.font.size
    return None

# --- helpers for summary tables ---
# --- helpers for summary overwrite ---
def _std_key(s: str) -> str:
    """
    Normalize a label to match a 'Std N' key.
    e.g. 'Std 7. Internal Review' -> 'std 7'
    """
    t = canon_label(s)
    m = re.match(r"(std\s+\d+)", t)
    return m.group(1) if m else t

def _looks_like_summary_table(table: Table) -> Optional[Tuple[int, int]]:
    """
    Return (label_col_idx, details_col_idx) if this is a Summary table
    with a DETAILS column; otherwise None.
    """
    if not table.rows:
        return None
    first = table.rows[0]
    cols = len(first.cells)
    if cols < 2:
        return None

    # header texts for first row
    head = [canon(cell_text(c)) for c in first.cells]

    # find DETAILS column
    details_col = None
    for j, t in enumerate(head):
        if "detail" in t:
            details_col = j
            break
    if details_col is None:
        return None

    # find the label column (left-hand standards column)
    label_col = None
    for j, t in enumerate(head):
        if any(k in t for k in ["maintenance management", "mass management", "fatigue management"]):
            label_col = j
            break
    if label_col is None:
        # fallback: assume the first non-DETAILS column is the label column
        label_col = 0 if details_col != 0 else 1

    return (label_col, details_col)
def count_header_rows(table: Table, scan_up_to: int = 6) -> int:
    """Heuristically count header rows (stop when first data row like '1.' appears)."""
    for i, row in enumerate(table.rows[:scan_up_to]):
        first = cell_text(row.cells[0]).strip()
        if re.match(r"^\d+\.?$", first):
            return i
    return 1
def _header_col_texts(table: Table, scan_rows: int = 5) -> List[str]:
    scan_rows = min(scan_rows, len(table.rows))
    if scan_rows == 0:
        return []
    # pick the row with the most cells as base
    base_row = max(range(scan_rows), key=lambda i: len(table.rows[i].cells))
    base_cols = len(table.rows[base_row].cells)
    cols = []
    for j in range(base_cols):
        parts = []
        for i in range(scan_rows):
            row = table.rows[i]
            if j < len(row.cells):
                parts.append(cell_text(row.cells[j]))
        cols.append(canon(" ".join(parts)))
    return cols

def count_header_rows(table: Table, scan_up_to: int = 6) -> int:
    """Header ends right before the first row whose 1st cell looks like '1.'"""
    limit = min(scan_up_to, len(table.rows))
    for i in range(limit):
        first = cell_text(table.rows[i].cells[0]).strip()
        if re.match(r"^\d+\.?$", first):
            return i
    # fallback to 1 header row
    return 1

def map_cols_mass_strict(table: Table) -> Dict[str, int]:
    cols = _header_col_texts(table, 5)
    def first_col(*needles):
        for j, t in enumerate(cols):
            if all(n in t for n in needles):
                return j
        return None
    idx = {
        "no":   first_col("no"),
        "reg":  first_col("registration", "number") or first_col("registration"),
        "wv":   first_col("weight", "verification"),
        "rfs":  first_col("rfs", "cert") or first_col("rfs", "certification"),
        "susp": first_col("suspension", "maintenance"),
        "trip": first_col("trip", "record"),
        "frs":  first_col("fault", "suspension") or first_col("fault", "reporting", "suspension"),
    }
    return {k: v for k, v in idx.items() if v is not None}

def find_mass_vehicle_numbers_table(doc: Document) -> Optional[Table]:
    """Pick the Mass vehicle-number table by matching its column set (not the Summary table)."""
    best = None
    best_score = -1
    for t in iter_tables(doc):
        cols = _header_col_texts(t, 5)
        allhdr = " ".join(cols)
        # must look like the vehicle numbers table
        hits = 0
        hits += int(any("registration" in c and "number" in c for c in cols))
        hits += int(any("weight" in c and "verification" in c for c in cols))
        hits += int(any("rfs" in c and ("cert" in c or "certification" in c) for c in cols))
        hits += int(any("suspension" in c and "maintenance" in c for c in cols))
        hits += int(any("trip" in c and "record" in c for c in cols))
        hits += int(any("fault" in c and "suspension" in c for c in cols))
        # reject obvious Summary tables
        if "details" in allhdr:
            continue
        # prefer tables with numbering column and many rows
        score = hits + (0.5 if any("no" == c or c.startswith("no ") for c in cols) else 0) + (len(t.rows) / 100.0)
        if hits >= 4 and score > best_score:
            best, best_score = t, score
    return best

def update_operator_declaration(doc: Document, print_name: str, position_title: str) -> bool:
    """
    First try strict table label mapping for 'Print Name' and 'Position Title'.
    If not found, fallback to the first two red placeholders under the 'Operator Declaration' heading.
    """
    changed = False
    # 1) Table label approach
    for lbl, val in (("Print Name", print_name), ("Position Title", position_title)):
        if not val:
            continue
        loc = find_label_cell(doc, lbl)
        if not loc:
            # tolerate odd spacing/colon/camelcase
            for alt in ("PrintName", "Print  Name", "Print Name:", "PositionTitle", "Position  Title", "Position Title:"):
                loc = find_label_cell(doc, alt)
                if loc:
                    break
        if loc:
            t, r, c = loc
            cell = get_adjacent_value_cell(t, r, c)
            if not replace_red_in_cell(cell, val):
                _set_cell_text_black(cell, val)
            changed = True

    if changed:
        return True

    # 2) Fallback: heading-scoped red placeholders
    head = "OPERATOR DECLARATION"
    p = find_heading_paragraph(doc, head) or find_heading_paragraph(doc, head.title())
    if not p:
        return False
    allp = iter_paragraphs(doc)
    try:
        i = allp.index(p)
    except ValueError:
        i = 0
    red_targets = []
    for q in allp[i+1:i+1+20]:
        reds = [r for r in q.runs if is_red_run(r)]
        if reds:
            red_targets.extend(reds)
        if len(red_targets) >= 2:
            break
    wrote = False
    if print_name and red_targets:
        _set_text_and_black(red_targets[0], print_name); wrote = True
    if position_title and len(red_targets) >= 2:
        _set_text_and_black(red_targets[1], position_title); wrote = True
    return wrote


def fill_mass_vehicle_table_preserve_headers(table: Table, arrays: Dict[str, List[str]]):
    colmap = map_cols_mass_strict(table)
    if "reg" not in colmap:
        return
    hdr_rows = count_header_rows(table, 6)
    regs = arrays.get("Registration Number", [])
    n = len(regs)

    # clear data rows only
    while len(table.rows) > hdr_rows:
        table._tbl.remove(table.rows[-1]._tr)
    # ensure enough rows
    while len(table.rows) < hdr_rows + n:
        table.add_row()

    def put(row, key, arr_key, i):
        if key in colmap:
            vals = arrays.get(arr_key, [])
            val = nz(vals[i]) if i < len(vals) else ""
            replace_red_in_cell(row.cells[colmap[key]], val)

    for i in range(n):
        row = table.rows[hdr_rows + i]
        replace_red_in_cell(row.cells[colmap["reg"]], nz(regs[i]))
        put(row, "wv",   "Weight Verification Records", i)
        put(row, "rfs",  "RFS Suspension Certification #", i)
        put(row, "susp", "Suspension System Maintenance", i)
        put(row, "trip", "Trip Records", i)
        put(row, "frs",  "Fault Recording/ Reporting on Suspension System", i)

def overwrite_summary_details_cells(doc: Document, section_name: str, section_dict: Dict[str, List[str]]) -> int:
    """For a Summary table (Maintenance/Mass/Fatigue), replace the entire DETAILS cell
    for each Std N row with the JSON text (written in black)."""
    # build desired texts
    desired: Dict[str, str] = { _std_key(k): join_value(v) for k, v in section_dict.items() }

    # pick which tables belong to this section by header sniff
    wanted_prefix = canon_label(section_name.split()[0])  # "maintenance" | "mass" | "fatigue"

    updated = 0
    for t in doc.tables:
        cols = _looks_like_summary_table(t)
        if not cols:
            continue
        label_col, details_col = cols

        head_txt = table_header_text(t, up_to_rows=2)
        if wanted_prefix not in head_txt:   # keep to the correct section
            continue

        # walk body rows
        for i in range(1, len(t.rows)):
            row = t.rows[i]
            key = _std_key(cell_text(row.cells[label_col]))

            # exact match or "std N" prefix match
            cand = desired.get(key)
            if not cand:
                m = re.match(r"(std\s+\d+)", key)
                if m:
                    for k2, v2 in desired.items():
                        if k2.startswith(m.group(1)):
                            cand = v2
                            break
            if not cand:
                continue

            _set_cell_text_black(row.cells[details_col], cand)  # full overwrite, black
            updated += 1
    return updated

SPLIT_SENT_PAT = re.compile(r"(?<=\.|\?|!)\s+")
ORDINAL_DATE_PAT = re.compile(r"\b(\d{1,2}(?:st|nd|rd|th)\s+[A-Za-z]+\s+\d{4})\b", re.I)

def split_sentences_keep(text: str) -> List[str]:
    s = " ".join(str(text or "").split())
    if not s:
        return []
    out = []
    start = 0
    for m in SPLIT_SENT_PAT.finditer(s):
        out.append(s[start:m.start()].strip())
        start = m.end()
    last = s[start:].strip()
    if last:
        out.append(last)
    return out

_sent_split = re.compile(r'(?<=[.!?])\s+|\n+')
_date_pat   = re.compile(r'\b(?:\d{1,2}(?:st|nd|rd|th)\s+[A-Za-z]+\s+\d{4}|\d{1,2}/\d{1,2}/\d{2,4}|[A-Za-z]+\s+\d{1,2},\s*\d{4})\b')

def extract_summary_snippets(desired_text: str):
    sents = _sentences(desired_text)
    dates = [m.group(0) for m in _date_pat.finditer(desired_text)]
    pick  = lambda rx: next((s for s in sents if re.search(rx, s, re.I)), None)
    return {
        "sheet_sent": pick(r'\b(daily\s+check|sheet)\b'),
        "sheet_phrase": _extract_sheet_phrase_from_desired(desired_text),
        "review":  pick(r'\binternal\s+review\b'),
        "qcs":     pick(r'\bquarterly\b.*\bcompliance\b') or pick(r'\bquarterly\b'),
        "dates":   dates,
        "sents":   sents,
    }

def fill_management_summary_tables(doc: Document, section_key: str, section_data: Dict[str, List[str]]):
    """
    Fill ALL summary tables for the given section_key ('maintenance'|'mass'|'fatigue')
    by matching each row label (left column) against keys in section_data and
    patching only the red text inside the DETAILS cell.
    """
    targets = [x for x in find_all_summary_tables(doc) if x[0] == section_key]
    if not targets:
        return

    # build list of (normalized label, original label, desired_text)
    desired = []
    for label, vals in section_data.items():
        want = canon_label(label)
        if not want:
            continue
        desired.append((want, label, join_value(vals)))

    for _, table, lcol, dcol in targets:
        # iterate data rows (skip header)
        for i in range(1, len(table.rows)):
            left_txt_norm = canon_label(cell_text(table.rows[i].cells[lcol]))
            if not left_txt_norm:
                continue
            for want_norm, _orig_lbl, value in desired:
                # loose contains match handles minor punctuation differences
                if want_norm and want_norm in left_txt_norm:
                    patch_details_cell_from_json(table.rows[i].cells[dcol], value)

def _set_text_and_black(run, new_text: str):
    """Replace a run's text and force color to black (clears theme color too)."""
    if new_text is None:
        new_text = ""
    run.text = str(new_text)
    run.font.color.rgb = BLACK
    try:
        # clear any theme color so rgb sticks
        run.font.color.theme_color = None
    except Exception:
        pass

def update_business_summary_once(doc: Document, value) -> bool:
    """Replace only the red summary paragraph; keep 'Accreditation Number' and 'Expiry Date' lines."""
    loc = (find_label_cell(doc, "Nature of the Operators Business (Summary)")
           or find_label_cell(doc, "Nature of the Operators Business (Summary):"))
    if not loc:
        return False

    t, r, c = loc
    cell = get_adjacent_value_cell(t, r, c)
    if not cell.paragraphs:
        cell.add_paragraph("")

    txt = join_value(value)

    # find paragraphs with any red runs (the placeholders for the summary)
    red_paras = [p for p in cell.paragraphs if any(is_red_run(run) for run in p.runs)]

    if red_paras:
        # write the summary into the first red paragraph (in black)
        _clear_para_and_write_black(red_paras[0], txt)
        # clear any extra red placeholders
        for p in red_paras[1:]:
            _clear_para_and_write_black(p, "")
    else:
        # no red placeholder found: just put the summary into the first paragraph, leave others
        _clear_para_and_write_black(cell.paragraphs[0], txt)

    return True


def _nuke_cell_paragraphs(cell: _Cell):
    """Remove ALL paragraphs from a cell (true delete, not just emptying runs)."""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

def _clear_para_and_write_black(paragraph, text: str):
    """Clear a whole paragraph and write fresh black text."""
    # wipe existing runs
    for r in list(paragraph.runs):
        r.text = ""
    r = paragraph.add_run(str(text or ""))
    r.font.color.rgb = BLACK
    try:
        r.font.color.theme_color = None
    except Exception:
        pass

def _set_cell_text_black(cell, text: str):
    """Clear a table cell and insert black text."""
    # remove text from all runs in all paragraphs
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    r = p.add_run(str(text or ""))
    r.font.color.rgb = BLACK
    try:
        r.font.color.theme_color = None
    except Exception:
        pass

def nz(x: Optional[str]) -> str:
    return (x or "").strip()

def canon(s: str) -> str:
    s = re.sub(r"\s+", " ", str(s)).strip().lower()
    s = s.replace("–", "-").replace("—", "-")
    return re.sub(r"[^a-z0-9/#()+,.\- ]+", "", s)

def canon_label(s: str) -> str:
    # labels often vary by punctuation/casing; keep digits/letters
    s = re.sub(r"\s+", " ", str(s)).strip().lower()
    s = s.replace("–", "-").replace("—", "-")
    s = re.sub(r"[^a-z0-9 ]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()

def join_value(value) -> str:
    if isinstance(value, list):
        # Keep multi-line when list provided
        return "\n".join([str(v) for v in value if nz(v)])
    return str(value)

def split_digits(s: str) -> List[str]:
    return re.findall(r"\d", s)

def para_text(p: Paragraph) -> str:
    return "".join(run.text for run in p.runs)

def cell_text(c: _Cell) -> str:
    return "\n".join(para_text(p) for p in c.paragraphs)

def is_red_run(run) -> bool:
    col = run.font.color
    if not col:
        return False
    if col.rgb is not None:
        return col.rgb == RED
    # Some templates use theme colors; treat explicit red text snippets only
    return False

def replace_red_in_paragraph(p: Paragraph, new_text: str) -> bool:
    replaced = False
    red_runs = [r for r in p.runs if is_red_run(r)]
    if not red_runs:
        return False
    # collapse all red runs into one and write value (in black)
    first = red_runs[0]
    _set_text_and_black(first, new_text)
    for r in red_runs[1:]:
        r.text = ""
    replaced = True
    return replaced

def replace_red_in_cell(cell: _Cell, new_text: str) -> bool:
    # replace only red runs; if none, replace whole cell with a single run (fallback)
    any_red = False
    for p in cell.paragraphs:
        if replace_red_in_paragraph(p, new_text):
            any_red = True
    if any_red:
        return True
    # fallback: clear cell, set single paragraph text in black
    _set_cell_text_black(cell, new_text)
    return True

def parse_attendance_lines(value) -> List[str]:
    """
    Parse strings like:
      "Peter Sheppard - Compliance Greg Dyer - Auditor"
    into:
      ["Peter Sheppard - Compliance", "Greg Dyer - Auditor"]
    Handles lists, newlines, semicolons, and pipes too.
    """
    if isinstance(value, list):
        s = " ".join(str(v) for v in value if v)
    else:
        s = str(value or "")
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return []

    # First split on explicit separators; then within each chunk, extract Name - Title pairs.
    chunks = re.split(r"\s*[\n;|]\s*", s)
    items: List[str] = []

    pair_pat = re.compile(
        r"([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){0,3})\s*-\s*"
        r"([^-\n]+?)(?=\s+[A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){0,3}\s*-\s*|$)"
    )

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        found = False
        for m in pair_pat.finditer(chunk):
            name = m.group(1).strip()
            title = m.group(2).strip()
            items.append(f"{name} - {title}")
            found = True
        if not found:
            # Fallback: single "Name - Title"
            if " - " in chunk:
                a, b = chunk.split(" - ", 1)
                items.append(f"{a.strip()} - {b.strip()}")
            elif chunk:
                items.append(chunk)

    return items

def fill_attendance_block(doc: Document, value) -> bool:
    items = parse_attendance_lines(value)
    if not items:
        return False

    loc = find_label_cell(doc, "Attendance List (Names and Position Titles)")
    if not loc:
        return False

    t, r, c = loc
    # value cell: usually directly under the heading cell
    target = (
        t.rows[r + 1].cells[c]
        if r + 1 < len(t.rows) and c < len(t.rows[r + 1].cells)
        else get_adjacent_value_cell(t, r, c)
    )

    # ---- read ONLY the target cell (don’t touch the row)
    def is_red_para(p): return any(is_red_run(run) for run in p.runs)
    def looks_like_pair(s: str) -> bool:
        if " - " not in s: return False
        a, b = s.split(" - ", 1)
        return bool(a.strip()) and bool(b.strip())

    paras = list(target.paragraphs)
    red_count = sum(1 for p in paras if is_red_para(p))
    existing_black = [para_text(p).strip() for p in paras
                      if (not is_red_para(p)) and looks_like_pair(para_text(p))]

    # compose final lines
    out_lines: List[str] = []
    out_lines.extend(items[:red_count])          # replace red placeholders
    out_lines.extend(existing_black)             # keep black lines
    norm = lambda s: re.sub(r"\s+", " ", s.strip().lower())
    seen = {norm(x) for x in out_lines}
    for extra in items[red_count:]:
        k = norm(extra)
        if k not in seen:
            out_lines.append(extra); seen.add(k)

    # ---- hard clear target cell and write fresh (all black)
    _nuke_cell_paragraphs(target)
    # first line
    p = target.add_paragraph()
    _clear_para_and_write_black(p, out_lines[0] if out_lines else "")
    # remaining lines
    for line in out_lines[1:]:
        p = target.add_paragraph()
        _clear_para_and_write_black(p, line)

    return True

# ----------------------------- document search -----------------------------
def iter_tables(doc: Document) -> List[Table]:
    return list(doc.tables)

def iter_paragraphs(doc: Document) -> List[Paragraph]:
    # paragraphs at doc level + inside tables
    out = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
    return out

def find_heading_paragraph(doc: Document, heading_text: str, window: int = 60) -> Optional[Paragraph]:
    key = canon(heading_text)
    for p in iter_paragraphs(doc):
        if canon(para_text(p)).startswith(key):
            return p
    # fuzzy contains
    for p in iter_paragraphs(doc):
        if key in canon(para_text(p)):
            return p
    return None

def find_label_cell_in_table(table: Table, label: str) -> Optional[Tuple[int, int]]:
    target = canon_label(label)
    for r_i, row in enumerate(table.rows):
        for c_i, cell in enumerate(row.cells):
            if canon_label(cell_text(cell)) == target:
                return (r_i, c_i)
    # allow contains (safe-ish)
    for r_i, row in enumerate(table.rows):
        for c_i, cell in enumerate(row.cells):
            if target and target in canon_label(cell_text(cell)):
                return (r_i, c_i)
    return None

def find_label_cell(doc: Document, label: str) -> Optional[Tuple[Table, int, int]]:
    for t in iter_tables(doc):
        pos = find_label_cell_in_table(t, label)
        if pos:
            return (t, pos[0], pos[1])
    return None

def get_adjacent_value_cell(table: Table, r: int, c: int) -> _Cell:
    # Prefer right cell, otherwise next row same col, otherwise this cell
    cols = len(table.rows[0].cells)
    if c + 1 < cols:
        return table.rows[r].cells[c+1]
    if r + 1 < len(table.rows):
        return table.rows[r+1].cells[c]
    return table.rows[r].cells[c]

# ----------------------------- label/value updates -----------------------------
def update_label_value_in_tables(doc: Document, label: str, value) -> bool:
    tup = find_label_cell(doc, label)
    val = join_value(value)
    if not tup:
        return False
    t, r, c = tup
    target_cell = get_adjacent_value_cell(t, r, c)
    return replace_red_in_cell(target_cell, val)

def update_heading_followed_red(doc: Document, heading: str, value, max_scan: int = 12) -> bool:
    """Find heading paragraph, then replace the first red run found within next N paragraphs (including inside tables)"""
    start = find_heading_paragraph(doc, heading)
    if not start:
        return False
    # Build a linear list of paragraphs across whole doc to get an index
    allp = iter_paragraphs(doc)
    try:
        idx = allp.index(start)
    except ValueError:
        idx = 0
    new_text = join_value(value)
    # Scan forward
    for p in allp[idx+1: idx+1+max_scan]:
        if replace_red_in_paragraph(p, new_text):
            return True
        # Also check any red in table cells inside this paragraph's parent (already covered via iter_paragraphs)
    return False

# ----------------------------- ACN per-digit fill -----------------------------
def fill_acn_digits(doc: Document, acn_value: str) -> bool:
    digits = split_digits(acn_value)
    if not digits:
        return False
    loc = find_label_cell(doc, "Australian Company Number")
    if not loc:
        return False

    t, r, c = loc

    # Collect cells to the RIGHT in the same row first
    targets: List[_Cell] = [t.rows[r].cells[j] for j in range(c + 1, len(t.rows[r].cells))]

    # If not enough, continue row-by-row below (left→right)
    rr = r + 1
    while len(targets) < len(digits) and rr < len(t.rows):
        targets.extend(list(t.rows[rr].cells))
        rr += 1

    targets = targets[:len(digits)]
    if not targets:
        return False

    # Clear each target cell and write ONE digit in black
    for d, cell in zip(digits, targets):
        _set_cell_text_black(cell, d)

    return True


# ----------------------------- vehicle tables -----------------------------
def table_header_text(table: Table, up_to_rows: int = 3) -> str:
    heads = []
    for i, row in enumerate(table.rows[:up_to_rows]):
        for cell in row.cells:
            heads.append(cell_text(cell))
    return canon(" ".join(heads))

def find_vehicle_table(doc: Document, want: str) -> Optional[Table]:
    """
    want = "maintenance" or "mass"
    """
    MAINT_KEYS = ["registration number", "maintenance records", "daily checks", "fault recording", "fault repair"]
    MASS_KEYS  = ["registration number", "weight verification", "rfs suspension", "suspension system maintenance", "trip records", "reporting on suspension"]
    candidates = []
    for t in iter_tables(doc):
        htxt = table_header_text(t)
        if want == "maintenance":
            if all(k in htxt for k in ["registration", "maintenance", "fault"]) and "suspension" not in htxt:
                candidates.append(t)
        elif want == "mass":
            if "suspension" in htxt and "weight" in htxt:
                candidates.append(t)
    # Prefer the one with most rows
    if not candidates:
        return None
    return max(candidates, key=lambda tb: len(tb.rows))

def map_cols(table: Table, want: str) -> Dict[str, int]:
    # map header columns by keywords from the first 2 rows that contain headers
    header_rows = table.rows[:2]
    col_texts = []
    cols = len(table.rows[0].cells)
    for j in range(cols):
        txt = " ".join(cell_text(r.cells[j]) for r in header_rows if j < len(r.cells))
        col_texts.append(canon(txt))
    idx = {}
    def first_col(*needles) -> Optional[int]:
        for j, t in enumerate(col_texts):
            if all(n in t for n in needles):
                return j
        return None
    if want == "maintenance":
        idx["reg"]   = first_col("registration")
        idx["rw"]    = first_col("roadworthiness")
        idx["mr"]    = first_col("maintenance", "records")
        idx["daily"] = first_col("daily", "check")
        idx["fr"]    = first_col("fault", "recording")
        idx["rep"]   = first_col("fault", "repair")
    else:
        idx["reg"]   = first_col("registration")
        idx["wv"]    = first_col("weight", "verification")
        idx["rfs"]   = first_col("rfs", "cert")
        idx["susp"]  = first_col("suspension", "maintenance")
        idx["trip"]  = first_col("trip", "record")
        idx["frs"]   = first_col("fault", "suspension")
    return {k:v for k,v in idx.items() if v is not None}

def clear_data_rows_keep_headers(table: Table, header_rows: int = 1):
    # Keep first header_rows, drop everything else
    while len(table.rows) > header_rows:
        table._tbl.remove(table.rows[-1]._tr)

def ensure_rows(table: Table, need_rows: int):
    # assumes 1 header row; add rows to reach need_rows + 1 total
    while len(table.rows) < need_rows + 1:
        table.add_row()

def fill_vehicle_table(table: Table, want: str, arrays: Dict[str, List[str]]):
    colmap = map_cols(table, want)
    if "reg" not in colmap:
        return
    if want == "maintenance":
        regs = arrays.get("Registration Number", [])
        rw   = arrays.get("Roadworthiness Certificates", [])
        mr   = arrays.get("Maintenance Records", [])
        daily= arrays.get("Daily Checks", [])
        fr   = arrays.get("Fault Recording/ Reporting", [])
        rep  = arrays.get("Fault Repair", [])
        n = len(regs)
        # keep header row(s), then fill N rows
        clear_data_rows_keep_headers(table, header_rows=1)
        ensure_rows(table, n)
        for i in range(n):
            row = table.rows[i+1]
            def put(col_key, vals):
                if col_key not in colmap or i >= len(vals): return
                c = row.cells[colmap[col_key]]
                replace_red_in_cell(c, nz(vals[i]))
            # write each col
            c_reg = row.cells[colmap["reg"]]; replace_red_in_cell(c_reg, nz(regs[i]))
            put("rw",   rw)
            put("mr",   mr)
            put("daily",daily)
            put("fr",   fr)
            put("rep",  rep)
    else:
        regs = arrays.get("Registration Number", [])
        wv   = arrays.get("Weight Verification Records", [])
        rfs  = arrays.get("RFS Suspension Certification #", [])
        susp = arrays.get("Suspension System Maintenance", [])
        trip = arrays.get("Trip Records", [])
        frs  = arrays.get("Fault Recording/ Reporting on Suspension System", [])
        n = len(regs)
        clear_data_rows_keep_headers(table, header_rows=1)
        ensure_rows(table, n)
        for i in range(n):
            row = table.rows[i+1]
            def put(col_key, vals):
                if col_key not in colmap or i >= len(vals): return
                c = row.cells[colmap[col_key]]
                replace_red_in_cell(c, nz(vals[i]))
            c_reg = row.cells[colmap["reg"]]; replace_red_in_cell(c_reg, nz(regs[i]))
            put("wv",   wv)
            put("rfs",  rfs)
            put("susp", susp)
            put("trip", trip)
            put("frs",  frs)

# ----------------------------- driver table -----------------------------
def find_driver_table(doc: Document) -> Optional[Table]:
    for t in iter_tables(doc):
        h = table_header_text(t)
        if "driver / scheduler" in h and ("fit for duty" in h or "work diary" in h):
            return t
    return None

def map_driver_cols(table: Table) -> Dict[str,int]:
    header_rows = table.rows[:2]
    cols = len(table.rows[0].cells)
    col_texts = []
    for j in range(cols):
        txt = " ".join(cell_text(r.cells[j]) for r in header_rows if j < len(r.cells))
        col_texts.append(canon(txt))
    idx = {}
    def first_col(*needles):
        for j, t in enumerate(col_texts):
            if all(n in t for n in needles):
                return j
        return None
    idx["name"]  = first_col("driver", "name")
    idx["roster"]= first_col("roster", "safe")
    idx["fit"]   = first_col("fit for duty")
    # Work diary might be split across two headers; match "work diary" OR "electronic work diary"
    wd = first_col("work diary") or first_col("electronic work diary")
    if wd is not None: idx["wd"] = wd
    return {k:v for k,v in idx.items() if v is not None}

def fill_driver_table(table: Table, arrays: Dict[str, List[str]]):
    colmap = map_driver_cols(table)
    if not colmap:
        return

    names   = arrays.get("Driver / Scheduler Name", [])
    rosters = arrays.get("Roster / Schedule / Safe Driving Plan (Date Range)", [])
    fit     = arrays.get("Fit for Duty Statement Completed (Yes/No)", [])
    wd      = arrays.get("Work Diary Pages (Page Numbers) Electronic Work Diary Records (Date Range)", [])

    n = max(len(rosters), len(fit), len(wd), len(names))
    clear_data_rows_keep_headers(table, header_rows=1)
    ensure_rows(table, n)

    has_any_name = any(str(x).strip() for x in names)

    for i in range(n):
        row = table.rows[i+1]
        if "name" in colmap and has_any_name:
            replace_red_in_cell(row.cells[colmap["name"]], names[i] if i < len(names) else "")
        if "roster" in colmap:
            replace_red_in_cell(row.cells[colmap["roster"]], rosters[i] if i < len(rosters) else "")
        if "fit" in colmap:
            replace_red_in_cell(row.cells[colmap["fit"]],     fit[i] if i < len(fit) else "")
        if "wd" in colmap:
            replace_red_in_cell(row.cells[colmap["wd"]],      wd[i]  if i < len(wd)  else "")



# ----------------------------- main mapping -----------------------------
def flatten_simple_sections(data: Dict) -> Dict[str, str]:
    """Collect simple label->single value mappings from top-level sections other than tables."""
    out = {}
    skip_sections = {
        "Vehicle Registration Numbers Maintenance",
        "Vehicle Registration Numbers Mass",
        "Driver / Scheduler Records Examined",
        "paragraphs",
        "Attendance List (Names and Position Titles)",
        "Nature of the Operators Business (Summary)",
        "Maintenance Management Summary",
        "Mass Management Summary",
        "Fatigue Management Summary",
    }
    for sec, kv in data.items():
        if sec in skip_sections: continue
        if not isinstance(kv, dict): continue
        for label, val in kv.items():
            out[f"{sec}::{label}"] = join_value(val)
    return out

def run(input_json: Path, template_docx: Path, output_docx: Path):
    with open(input_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(str(template_docx))

    # 1) simple label/value tables
    simple = flatten_simple_sections(data)

    # Map by (section::label). We try: (a) find exact label cell somewhere and write in the adjacent cell;
    # (b) if not found, search by heading then the next red run below the heading.
    for k, v in simple.items():
        # use the part after '::' as the label
        label = k.split("::", 1)[1] if "::" in k else k

        # SPECIAL: skip ACN here; we'll fill per-digit later
        if canon_label(label) == "australian company number":
            continue

        ok = update_label_value_in_tables(doc, label, v)
        if not ok:
            sec = k.split("::", 1)[0] if "::" in k else k
            update_heading_followed_red(doc, sec, v)


            # 2) paragraphs block
        paras = data.get("paragraphs", {})

        # 2a) generic headings → replace next red (skip the 3 management headings here)
        # third-line headings above the three tables
        for head in ("MAINTENANCE MANAGEMENT", "MASS MANAGEMENT", "FATIGUE MANAGEMENT"):
            name_val = join_value(paras.get(head, ""))
            if name_val:
                update_heading_followed_red(doc, head, name_val, max_scan=6)

        # 2b) the 3-layer headings → overwrite the 3rd line only
        # second-last page: date under page heading
        aud_head = "NHVAS APPROVED AUDITOR DECLARATION"
        aud_date = join_value(paras.get(aud_head, ""))
        if aud_date:
            set_date_by_heading_from_end(doc, aud_head, aud_date, max_scan=40)

        # last page: date under the long acknowledgement paragraph
        ack_head = ("I hereby acknowledge and agree with the findings detailed in this NHVAS Audit Summary Report. "
                    "I have read and understand the conditions applicable to the Scheme, including the NHVAS Business Rules and Standards.")
        ack_date = join_value(paras.get(ack_head, ""))
        if ack_date:
            set_date_by_paragraph_from_end(doc, ack_head, ack_date, max_scan=40)

        maint_name = join_value(paras.get("MAINTENANCE MANAGEMENT", ""))
        if maint_name:
            set_layer3_name_after_management_heading(
                doc,
                "MAINTENANCE MANAGEMENT",
                ["Vehicle Registration Numbers of Records Examined"],
                maint_name,
            )

        mass_name = join_value(paras.get("MASS MANAGEMENT", ""))
        if mass_name:
            set_layer3_name_after_management_heading(
                doc,
                "MASS MANAGEMENT",
                ["Vehicle Registration Numbers of Records Examined"],
                mass_name,
            )

        fat_name = join_value(paras.get("FATIGUE MANAGEMENT", ""))
        if fat_name:
            set_layer3_name_after_management_heading(
                doc,
                "FATIGUE MANAGEMENT",
                ["Driver / Scheduler Records Examined"],
                fat_name,
            )


    # 3) ACN digits
    op_info = data.get("Operator Information", {})
    acn_val = join_value(op_info.get("Australian Company Number", ""))
    if acn_val:
        fill_acn_digits(doc, acn_val)

    # 4) Vehicle tables
    maint = data.get("Vehicle Registration Numbers Maintenance", {})
    mass  = data.get("Vehicle Registration Numbers Mass", {})
    t_m = find_vehicle_table(doc, "maintenance")
    if t_m and maint:
        fill_vehicle_table(t_m, "maintenance", maint)
    t_ms = find_mass_vehicle_numbers_table(doc)
    if t_ms and mass:
        fill_mass_vehicle_table_preserve_headers(t_ms, mass)

    # 5) Driver table
    drivers = data.get("Driver / Scheduler Records Examined", {})
    t_d = find_driver_table(doc)
    if t_d and drivers:
        fill_driver_table(t_d, drivers)

    # 6) Special: Audit Declaration dates via heading
    decl = data.get("Audit Declaration dates", {})
    if decl.get("Audit was conducted on"):
        update_heading_followed_red(doc, "Audit was conducted on", decl["Audit was conducted on"])

    # 7) Operator Declaration (last page, bottom row only), and fix Auditor table header
    op_decl = data.get("Operator Declaration", {})
    if op_decl:
        fill_operator_declaration(
            doc,
            join_value(op_decl.get("Print Name", "")),
            join_value(op_decl.get("Position Title", "")),
        )

    # make sure the second-last page “NHVAS APPROVED AUDITOR DECLARATION” header row is labels
    ensure_auditor_decl_headers(doc)


    # 8) Attendance List
    # Attendance: replace red lines only
    atts = data.get("Attendance List (Names and Position Titles)", {})
    att_val = atts.get("Attendance List (Names and Position Titles)")
    if att_val:
        fill_attendance_block(doc, att_val)

    # 9) Nature of the Operators Business (Summary): write once (no duplicates)
    biz = data.get("Nature of the Operators Business (Summary)", {})
    if biz:
        val = biz.get("Nature of the Operators Business (Summary):") or next(iter(biz.values()), "")
        if val:
            update_business_summary_once(doc, val)

    # 10) Summary tables: FULL OVERWRITE of DETAILS from JSON
    mm_sum = data.get("Maintenance Management Summary", {})
    if mm_sum:
        overwrite_summary_details_cells(doc, "Maintenance Management Summary", mm_sum)

    mass_sum = data.get("Mass Management Summary", {})
    if mass_sum:
        overwrite_summary_details_cells(doc, "Mass Management Summary", mass_sum)

    fat_sum = data.get("Fatigue Management Summary", {})
    if fat_sum:
        overwrite_summary_details_cells(doc, "Fatigue Management Summary", fat_sum)


    doc.save(str(output_docx))

# ----------------------------- CLI -----------------------------
if __name__ == "__main__":
    import sys
    from pathlib import Path

    if len(sys.argv) != 4:
        print("Usage: python updated_word.py <json> <template.docx> <output.docx>")
        sys.exit(1)

    a, b, c = map(Path, sys.argv[1:4])
    files = [a, b, c]

    json_path = next((p for p in files if p.suffix.lower() == ".json"), None)
    docx_paths = [p for p in files if p.suffix.lower() == ".docx"]

    if not json_path or len(docx_paths) < 2:
        print("Error: provide one .json and two .docx (template + output).")
        sys.exit(1)

    # Template = the .docx that already exists; Output = the other .docx
    template_docx = next((p for p in docx_paths if p.exists()), docx_paths[0])
    output_docx = docx_paths[1] if docx_paths[0] == template_docx else docx_paths[0]

    run(json_path, template_docx, output_docx)