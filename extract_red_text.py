#!/usr/bin/env python3
import re
import json
import sys
from docx import Document
from docx.oxml.ns import qn
from master_key import TABLE_SCHEMAS, HEADING_PATTERNS, PARAGRAPH_PATTERNS

def normalize_header_label(s: str) -> str:
    """Normalize a header/label by stripping parentheticals & punctuation."""
    s = re.sub(r"\s+", " ", s.strip())
    # remove content in parentheses/brackets
    s = re.sub(r"\([^)]*\)", "", s)
    s = re.sub(r"\[[^]]*\]", "", s)
    # unify slashes and hyphens, collapse spaces
    s = s.replace("–", "-").replace("—", "-").replace("/", " / ").replace("  ", " ")
    return s.strip()

# Canonical label aliases for Vehicle/Maintenance/General headers
LABEL_ALIASES = {
    # Vehicle Registration (Maintenance)
    "roadworthiness certificates": "Roadworthiness Certificates",
    "maintenance records": "Maintenance Records",
    "daily checks": "Daily Checks",
    "fault recording / reporting": "Fault Recording/ Reporting",
    "fault repair": "Fault Repair",

    # Vehicle Registration (Mass)
    "sub contracted vehicles statement of compliance": "Sub-contracted Vehicles Statement of Compliance",
    "weight verification records": "Weight Verification Records",
    "rfs suspension certification #": "RFS Suspension Certification #",
    "suspension system maintenance": "Suspension System Maintenance",
    "trip records": "Trip Records",
    "fault recording/ reporting on suspension system": "Fault Recording/ Reporting on Suspension System",

    # Common
    "registration number": "Registration Number",
    "no.": "No.",
    "sub contractor": "Sub contractor",
    "sub-contractor": "Sub contractor",
}

def looks_like_operator_declaration(context):
    """True iff heading says Operator Declaration and headers include Print Name + Position Title."""
    heading = (context.get("heading") or "").strip().lower()
    headers = " ".join(context.get("headers") or []).lower()
    return (
        "operator declaration" in heading
        and "print name" in headers
        and "position" in headers
        and "title" in headers
    )

def looks_like_auditor_declaration(context):
    heading = (context.get("heading") or "").strip().lower()
    headers = " ".join(context.get("headers") or []).lower()
    return (
        "auditor declaration" in heading
        and "print name" in headers
        and ("nhvr" in headers or "auditor registration number" in headers)
    )

# --- NEW: header-only fallback that ignores headings and just keys on the two column names
def extract_operator_declaration_by_headers_from_end(doc):
    """
    Scan tables from the end; if a table's first row contains both
    'Print Name' AND 'Position Title' (case-insensitive), extract red text
    from the data rows into:
        {"Print Name": [...], "Position Title": [...]}
    """
    for tbl in reversed(doc.tables):
        if len(tbl.rows) < 2:
            continue  # need header + at least one data row

        headers_norm = [normalize_header_label(c.text).lower() for c in tbl.rows[0].cells]
        has_print   = any("print name" in h for h in headers_norm)
        has_pos_tit = any(("position title" in h) or ("position" in h and "title" in h) for h in headers_norm)
        if not (has_print and has_pos_tit):
            continue

        idx_print = next((i for i, h in enumerate(headers_norm) if "print name" in h), None)
        idx_pos   = next((i for i, h in enumerate(headers_norm) if "position title" in h), None)
        if idx_pos is None:
            idx_pos = next((i for i, h in enumerate(headers_norm) if ("position" in h and "title" in h)), None)

        result = {"Print Name": [], "Position Title": []}
        for row in tbl.rows[1:]:
            if idx_print is not None and idx_print < len(row.cells):
                cell = row.cells[idx_print]
                reds = [r.text for p in cell.paragraphs for r in p.runs if is_red_font(r) and r.text]
                reds = coalesce_numeric_runs(reds)
                txt  = normalize_text(" ".join(reds))
                if txt:
                    result["Print Name"].append(txt)

            if idx_pos is not None and idx_pos < len(row.cells):
                cell = row.cells[idx_pos]
                reds = [r.text for p in cell.paragraphs for r in p.runs if is_red_font(r) and r.text]
                reds = coalesce_numeric_runs(reds)
                txt  = normalize_text(" ".join(reds))
                if txt:
                    result["Position Title"].append(txt)

        if result["Print Name"] or result["Position Title"]:
            return {k: v for k, v in result.items() if v}

    return None
# --- end NEW helper

def canonicalize_label(s: str) -> str:
    key = normalize_header_label(s).lower()
    key = re.sub(r"\s+", " ", key)
    return LABEL_ALIASES.get(key, s)

def bag_similarity(a: str, b: str) -> float:
    """Loose bag-of-words similarity for header↔label matching."""
    aw = {w for w in re.split(r"[^A-Za-z0-9#]+", normalize_header_label(a).lower()) if len(w) > 2 or w in {"#","no"}}
    bw = {w for w in re.split(r"[^A-Za-z0-9#]+", normalize_header_label(b).lower()) if len(w) > 2 or w in {"#","no"}}
    if not aw or not bw: 
        return 0.0
    inter = len(aw & bw)
    return inter / max(len(aw), len(bw))

def coalesce_numeric_runs(text_list):
    """
    If a cell yields ['4','5','6','9','8','7','1','2','3'] etc., join continuous single-char digit runs.
    Returns ['456987123'] instead of many singles. Non-digit tokens are preserved.
    """
    out, buf = [], []
    for t in text_list:
        if len(t) == 1 and t.isdigit():
            buf.append(t)
        else:
            if buf:
                out.append("".join(buf))
                buf = []
            out.append(t)
    if buf:
        out.append("".join(buf))
    return out

def is_red_font(run):
    """Enhanced red font detection with better color checking"""
    col = run.font.color
    if col and col.rgb:
        r, g, b = col.rgb
        if r > 150 and g < 100 and b < 100 and (r-g) > 30 and (r-b) > 30:
            return True
    rPr = getattr(run._element, "rPr", None)
    if rPr is not None:
        clr = rPr.find(qn('w:color'))
        if clr is not None:
            val = clr.get(qn('w:val'))
            if val and re.fullmatch(r"[0-9A-Fa-f]{6}", val):
                rr, gg, bb = int(val[:2], 16), int(val[2:4], 16), int(val[4:], 16)
                if rr > 150 and gg < 100 and bb < 100 and (rr-gg) > 30 and (rr-bb) > 30:
                    return True
    return False

def _prev_para_text(tbl):
    """Get text from previous paragraph before table"""
    prev = tbl._tbl.getprevious()
    while prev is not None and not prev.tag.endswith("}p"):
        prev = prev.getprevious()
    if prev is None:
        return ""
    return "".join(node.text for node in prev.iter() if node.tag.endswith("}t") and node.text).strip()

def normalize_text(text):
    """Normalize text for better matching"""
    return re.sub(r'\s+', ' ', text.strip())

def fuzzy_match_heading(heading, patterns):
    """Check if heading matches any pattern with fuzzy matching"""
    heading_norm = normalize_text(heading.upper())
    for pattern in patterns:
        if re.search(pattern, heading_norm, re.IGNORECASE):
            return True
    return False

def get_table_context(tbl):
    """Get comprehensive context information for table"""
    heading = normalize_text(_prev_para_text(tbl))
    headers = [normalize_text(c.text) for c in tbl.rows[0].cells if c.text.strip()]
    col0 = [normalize_text(r.cells[0].text) for r in tbl.rows if r.cells[0].text.strip()]
    first_cell = normalize_text(tbl.rows[0].cells[0].text) if tbl.rows else ""
    all_cells = []
    for row in tbl.rows:
        for cell in row.cells:
            text = normalize_text(cell.text)
            if text:
                all_cells.append(text)
    return {
        'heading': heading,
        'headers': headers,
        'col0': col0,
        'first_cell': first_cell,
        'all_cells': all_cells,
        'num_rows': len(tbl.rows),
        'num_cols': len(tbl.rows[0].cells) if tbl.rows else 0
    }

def calculate_schema_match_score(schema_name, spec, context):
    """Enhanced calculate match score - IMPROVED for Vehicle Registration tables"""
    score = 0
    reasons = []
    
    # 🎯 VEHICLE REGISTRATION BOOST
    if "Vehicle Registration" in schema_name:
        vehicle_keywords = ["registration", "vehicle", "sub-contractor", "weight verification", "rfs suspension"]
        table_text = " ".join(context['headers']).lower() + " " + context['heading'].lower()
        keyword_matches = sum(1 for keyword in vehicle_keywords if keyword in table_text)
        if keyword_matches >= 2:
            score += 150  # Very high boost for vehicle tables
            reasons.append(f"Vehicle Registration keywords: {keyword_matches}/5")
        elif keyword_matches >= 1:
            score += 75   # Medium boost
            reasons.append(f"Some Vehicle Registration keywords: {keyword_matches}/5")
    
    # 🎯 SUMMARY TABLE BOOST (existing logic)
    if "Summary" in schema_name and "details" in " ".join(context['headers']).lower():
        score += 100
        reasons.append(f"Summary schema with DETAILS column - perfect match")
    
    if "Summary" not in schema_name and "details" in " ".join(context['headers']).lower():
        score -= 75
        reasons.append(f"Non-summary schema penalized for DETAILS column presence")
    
    # Context exclusions
    if spec.get("context_exclusions"):
        table_text = " ".join(context['headers']).lower() + " " + context['heading'].lower()
        for exclusion in spec["context_exclusions"]:
            if exclusion.lower() in table_text:
                score -= 50
                reasons.append(f"Context exclusion penalty: '{exclusion}' found")
    
    # Context keywords
    if spec.get("context_keywords"):
        table_text = " ".join(context['headers']).lower() + " " + context['heading'].lower()
        keyword_matches = 0
        for keyword in spec["context_keywords"]:
            if keyword.lower() in table_text:
                keyword_matches += 1
        
        if keyword_matches > 0:
            score += keyword_matches * 15
            reasons.append(f"Context keyword matches: {keyword_matches}/{len(spec['context_keywords'])}")
    
    # Direct first cell match
    if context['first_cell'] and context['first_cell'].upper() == schema_name.upper():
        score += 100
        reasons.append(f"Direct first cell match: '{context['first_cell']}'")
    
    # Heading pattern matching
    if spec.get("headings"):
        for h in spec["headings"]:
            if fuzzy_match_heading(context['heading'], [h["text"]]):
                score += 50
                reasons.append(f"Heading match: '{context['heading']}'")
                break
    
    # Column header matching
    if spec.get("columns"):
        cols = [normalize_text(col) for col in spec["columns"]]
        matches = 0
        for col in cols:
            if any(col.upper() in h.upper() for h in context['headers']):
                matches += 1
        if matches == len(cols):
            score += 60
            reasons.append(f"All column headers match: {cols}")
        elif matches > 0:
            score += matches * 20
            reasons.append(f"Partial column matches: {matches}/{len(cols)}")
    
    # Label matching for left-oriented tables
    if spec.get("orientation") == "left":
        labels = [normalize_text(lbl) for lbl in spec["labels"]]
        matches = 0
        for lbl in labels:
            if any(lbl.upper() in c.upper() or c.upper() in lbl.upper() for c in context['col0']):
                matches += 1
        if matches > 0:
            score += (matches / len(labels)) * 30
            reasons.append(f"Left orientation label matches: {matches}/{len(labels)}")
    
    # 🎯 ENHANCED Label matching for row1-oriented tables (Vehicle Registration)
    elif spec.get("orientation") == "row1":
        labels = [normalize_text(lbl) for lbl in spec["labels"]]
        matches = 0
        for lbl in labels:
            if any(lbl.upper() in h.upper() or h.upper() in lbl.upper() for h in context['headers']):
                matches += 1
            elif any(word.upper() in " ".join(context['headers']).upper() for word in lbl.split() if len(word) > 3):
                matches += 0.5  # Partial credit
        if matches > 0:
            score += (matches / len(labels)) * 40
            reasons.append(f"Row1 orientation header matches: {matches}/{len(labels)}")
    
    # Special handling for Declaration tables (existing logic)
    if schema_name == "Operator Declaration" and context['first_cell'].upper() == "PRINT NAME":
        if "OPERATOR DECLARATION" in context['heading'].upper():
            score += 80
            reasons.append("Operator Declaration context match")
        elif any("MANAGER" in cell.upper() for cell in context['all_cells']):
            score += 60
            reasons.append("Manager found in cells (likely Operator Declaration)")
    
    if schema_name == "NHVAS Approved Auditor Declaration" and context['first_cell'].upper() == "PRINT NAME":
        if any("MANAGER" in cell.upper() for cell in context['all_cells']):
            score -= 50
            reasons.append("Penalty: Manager found (not auditor)")
    
    return score, reasons

def match_table_schema(tbl):
    """Improved table schema matching with scoring system"""
    context = get_table_context(tbl)
    # Auditor Declaration first
    if ("print name" in " ".join(context.get("headers", [])).lower() and
        "auditor" in " ".join(context.get("headers", [])).lower()):
        return "NHVAS Approved Auditor Declaration"
    # NEW: prioritize Auditor Declaration to avoid misclassification
    if looks_like_auditor_declaration(context):
        return "NHVAS Approved Auditor Declaration"
    # hard-match Operator Declaration first (high priority, avoids misclassification)
    if looks_like_operator_declaration(context):
        return "Operator Declaration"
    best_match = None
    best_score = 0
    for name, spec in TABLE_SCHEMAS.items():
        score, reasons = calculate_schema_match_score(name, spec, context)
        if score > best_score:
            best_score = score
            best_match = name
    if best_score >= 20:
        return best_match
    return None

def check_multi_schema_table(tbl):
    """Check if table contains multiple schemas and split appropriately"""
    context = get_table_context(tbl)
    operator_labels = ["Operator name (Legal entity)", "NHVAS Accreditation No.", "Registered trading name/s", 
                      "Australian Company Number", "NHVAS Manual"]
    contact_labels = ["Operator business address", "Operator Postal address", "Email address", "Operator Telephone Number"]
    has_operator = any(any(op_lbl.upper() in cell.upper() for op_lbl in operator_labels) for cell in context['col0'])
    has_contact = any(any(cont_lbl.upper() in cell.upper() for cont_lbl in contact_labels) for cell in context['col0'])
    if has_operator and has_contact:
        return ["Operator Information", "Operator contact details"]
    return None

def extract_multi_schema_table(tbl, schemas):
    """Extract data from table with multiple schemas"""
    result = {}
    for schema_name in schemas:
        if schema_name not in TABLE_SCHEMAS:
            continue
        spec = TABLE_SCHEMAS[schema_name]
        schema_data = {}
        for ri, row in enumerate(tbl.rows):
            if ri == 0:
                continue
            row_label = normalize_text(row.cells[0].text)
            belongs_to_schema = False
            matched_label = None
            for spec_label in spec["labels"]:
                spec_norm = normalize_text(spec_label).upper()
                row_norm = row_label.upper()
                if spec_norm == row_norm or spec_norm in row_norm or row_norm in spec_norm:
                    belongs_to_schema = True
                    matched_label = spec_label
                    break
            if not belongs_to_schema:
                continue
            for ci, cell in enumerate(row.cells):
                red_txt = "".join(run.text for p in cell.paragraphs for run in p.runs if is_red_font(run)).strip()
                if red_txt:
                    if matched_label not in schema_data:
                        schema_data[matched_label] = []
                    if red_txt not in schema_data[matched_label]:
                        schema_data[matched_label].append(red_txt)
        if schema_data:
            result[schema_name] = schema_data
    return result

def extract_table_data(tbl, schema_name, spec):
    """Extract red text data from table based on schema – per-row repeats for specific tables."""

    # ───────────────────────────────────────────────────────────────────────────
    # OPERATOR DECLARATION (row1 headers: Print Name | Position Title)
    # ───────────────────────────────────────────────────────────────────────────
    if schema_name == "Operator Declaration":
        print(f"    🧾 EXTRACTION FIX: Processing Operator Declaration table")

        labels = spec["labels"]  # ["Print Name", "Position Title"]
        canonical_labels = {canonicalize_label(lbl): lbl for lbl in labels}

        collected = {lbl: [] for lbl in labels}

        if len(tbl.rows) < 2:
            print(f"    ❌ Operator Declaration table has less than 2 rows")
            return {}

        # map header cells → labels (row1 orientation)
        header_row = tbl.rows[0]
        column_mapping = {}
        print(f"    📋 Mapping {len(header_row.cells)} header cells to labels")

        for col_idx, cell in enumerate(header_row.cells):
            raw_h = normalize_text(cell.text)
            header_text = normalize_header_label(raw_h)
            if not header_text:
                continue
            print(f"      Column {col_idx}: '{raw_h}'")

            # alias/canonical first
            canon = canonicalize_label(header_text)
            if canon in canonical_labels:
                best_label = canonical_labels[canon]
                print(f"        ✅ Mapped to: '{best_label}' (alias)")
                column_mapping[col_idx] = best_label
                continue

            # else bag-of-words similarity
            best_label, best_score = None, 0.0
            for canon_lab, original_lab in canonical_labels.items():
                s = bag_similarity(header_text, canon_lab)
                if s > best_score:
                    best_score, best_label = s, original_lab

            if best_label and best_score >= 0.40:
                print(f"        ✅ Mapped to: '{best_label}' (score: {best_score:.2f})")
                column_mapping[col_idx] = best_label
            else:
                print(f"        ⚠️ No mapping found for '{raw_h}'")

        print(f"    📊 Total column mappings: {len(column_mapping)}")

        # collect red text from the (usually single) data row
        for row_idx in range(1, len(tbl.rows)):
            row = tbl.rows[row_idx]
            print(f"      📌 Processing data row {row_idx}")
            for col_idx, cell in enumerate(row.cells):
                if col_idx not in column_mapping:
                    continue
                label = column_mapping[col_idx]
                reds = [run.text for p in cell.paragraphs for run in p.runs if is_red_font(run) and run.text]
                if not reds:
                    continue
                reds = coalesce_numeric_runs(reds)
                red_txt = normalize_text(" ".join(reds))
                if not red_txt:
                    continue
                print(f"        🔴 Found red text in '{label}': '{red_txt}'")
                collected[label].append(red_txt)

        result = {k: v for k, v in collected.items() if v}
        print(f"    ✅ Operator Declaration extracted: {len(result)} columns with data")
        return result

    # ───────────────────────────────────────────────────────────────────────────
    # A) Vehicle Registration tables (per-row accumulation; NO dedupe)
    # ───────────────────────────────────────────────────────────────────────────
    if "Vehicle Registration" in schema_name:
        print(f"    🚗 EXTRACTION FIX: Processing Vehicle Registration table")

        labels = spec["labels"]
        canonical_labels = {canonicalize_label(lbl): lbl for lbl in labels}

        collected = {lbl: [] for lbl in labels}   # ← keep every row value
        unmapped_bucket = {}

        if len(tbl.rows) < 2:
            print(f"    ❌ Vehicle table has less than 2 rows")
            return {}

        header_row = tbl.rows[0]
        column_mapping = {}
        print(f"    📋 Mapping {len(header_row.cells)} header cells to labels")

        for col_idx, cell in enumerate(header_row.cells):
            raw_h = normalize_text(cell.text)
            header_text = normalize_header_label(raw_h)
            if not header_text:
                continue
            print(f"      Column {col_idx}: '{raw_h}'")

            # Try alias/canonical first
            canon = canonicalize_label(header_text)
            if canon in canonical_labels:
                best_label = canonical_labels[canon]
                print(f"        ✅ Mapped to: '{best_label}' (alias)")
                column_mapping[col_idx] = best_label
                continue

            # Else bag-of-words similarity
            best_label, best_score = None, 0.0
            for canon_lab, original_lab in canonical_labels.items():
                s = bag_similarity(header_text, canon_lab)
                if s > best_score:
                    best_score, best_label = s, original_lab

            if best_label and best_score >= 0.40:
                print(f"        ✅ Mapped to: '{best_label}' (score: {best_score:.2f})")
                column_mapping[col_idx] = best_label
            else:
                print(f"        ⚠️ No mapping found for '{raw_h}'")
                unmapped_bucket[raw_h] = []

        print(f"    📊 Total column mappings: {len(column_mapping)}")

        header_texts = [normalize_text(hc.text) for hc in header_row.cells]
        for row_idx in range(1, len(tbl.rows)):
            row = tbl.rows[row_idx]
            print(f"      📌 Processing data row {row_idx}")
            for col_idx, cell in enumerate(row.cells):
                reds = [run.text for p in cell.paragraphs for run in p.runs if is_red_font(run) and run.text]
                if not reds:
                    continue
                reds = coalesce_numeric_runs(reds)
                red_txt = normalize_text(" ".join(reds))
                if not red_txt:
                    continue

                if col_idx in column_mapping:
                    label = column_mapping[col_idx]
                    print(f"        🔴 Found red text in '{label}': '{red_txt}'")
                    collected[label].append(red_txt)  # ← append every occurrence
                else:
                    header_name = header_texts[col_idx] if col_idx < len(header_texts) else f"(unmapped col {col_idx})"
                    unmapped_bucket.setdefault(header_name, []).append(red_txt)

        result = {k: v for k, v in collected.items() if v}
        if unmapped_bucket:
            result.update({f"UNMAPPED::{k}": v for k, v in unmapped_bucket.items() if v})
        print(f"    ✅ Vehicle Registration extracted: {len(result)} columns with data")
        return result

    # ───────────────────────────────────────────────────────────────────────────
    # B) Driver / Scheduler Records Examined (per-row accumulation; NO dedupe)
    # ───────────────────────────────────────────────────────────────────────────
    if "Driver / Scheduler" in schema_name:
        print(f"    👤 EXTRACTION FIX: Processing Driver / Scheduler table")

        labels = spec["labels"]
        canonical_labels = {canonicalize_label(lbl): lbl for lbl in labels}

        collected = {lbl: [] for lbl in labels}   # ← keep every row value
        unmapped_bucket = {}

        if len(tbl.rows) < 2:
            print(f"    ❌ Driver/Scheduler table has less than 2 rows")
            return {}

        header_row = tbl.rows[0]
        column_mapping = {}
        print(f"    📋 Mapping {len(header_row.cells)} header cells to labels")

        for col_idx, cell in enumerate(header_row.cells):
            raw_h = normalize_text(cell.text)
            header_text = normalize_header_label(raw_h)
            if not header_text:
                continue
            print(f"      Column {col_idx}: '{raw_h}'")

            # Try alias/canonical first (rarely used here, but safe)
            canon = canonicalize_label(header_text)
            if canon in canonical_labels:
                best_label = canonical_labels[canon]
                print(f"        ✅ Mapped to: '{best_label}' (alias)")
                column_mapping[col_idx] = best_label
                continue

            # Else bag-of-words similarity (good for long headings)
            best_label, best_score = None, 0.0
            for canon_lab, original_lab in canonical_labels.items():
                s = bag_similarity(header_text, canon_lab)
                if s > best_score:
                    best_score, best_label = s, original_lab

            if best_label and best_score >= 0.40:
                print(f"        ✅ Mapped to: '{best_label}' (score: {best_score:.2f})")
                column_mapping[col_idx] = best_label
            else:
                print(f"        ⚠️ No mapping found for '{raw_h}'")
                unmapped_bucket[raw_h] = []

        print(f"    📊 Total column mappings: {len(column_mapping)}")

        header_texts = [normalize_text(hc.text) for hc in header_row.cells]
        for row_idx in range(1, len(tbl.rows)):
            row = tbl.rows[row_idx]
            print(f"      📌 Processing data row {row_idx}")
            for col_idx, cell in enumerate(row.cells):
                reds = [run.text for p in cell.paragraphs for run in p.runs if is_red_font(run) and run.text]
                if not reds:
                    continue
                reds = coalesce_numeric_runs(reds)
                red_txt = normalize_text(" ".join(reds))
                if not red_txt:
                    continue

                if col_idx in column_mapping:
                    label = column_mapping[col_idx]
                    print(f"        🔴 Found red text in '{label}': '{red_txt}'")
                    collected[label].append(red_txt)  # ← append every occurrence
                else:
                    header_name = header_texts[col_idx] if col_idx < len(header_texts) else f"(unmapped col {col_idx})"
                    unmapped_bucket.setdefault(header_name, []).append(red_txt)

        result = {k: v for k, v in collected.items() if v}
        if unmapped_bucket:
            result.update({f"UNMAPPED::{k}": v for k, v in unmapped_bucket.items() if v})
        print(f"    ✅ Driver / Scheduler extracted: {len(result)} columns with data")
        return result

    # ───────────────────────────────────────────────────────────────────────────
    # C) Generic tables (unchanged: WITH dedupe)
    # ───────────────────────────────────────────────────────────────────────────
    labels = spec["labels"] + [schema_name]
    collected = {lbl: [] for lbl in labels}
    seen = {lbl: set() for lbl in labels}
    by_col = (spec.get("orientation") == "row1")
    start_row = 1 if by_col else 0
    rows = tbl.rows[start_row:]

    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row.cells):
            reds = [run.text for p in cell.paragraphs for run in p.runs if is_red_font(run) and run.text]
            if not reds:
                continue
            reds = coalesce_numeric_runs(reds)
            red_txt = normalize_text(" ".join(reds))
            if not red_txt:
                continue

            if by_col:
                if ci < len(spec["labels"]):
                    lbl = spec["labels"][ci]
                else:
                    lbl = schema_name
            else:
                raw_label = normalize_text(row.cells[0].text)
                lbl = None
                for spec_label in spec["labels"]:
                    if normalize_text(spec_label).upper() == raw_label.upper():
                        lbl = spec_label
                        break
                if not lbl:
                    a_raw = normalize_header_label(raw_label).upper()
                    for spec_label in spec["labels"]:
                        a_spec = normalize_header_label(spec_label).upper()
                        if a_spec in a_raw or a_raw in a_spec:
                            lbl = spec_label
                            break
                if not lbl:
                    lbl = schema_name

            if red_txt not in seen[lbl]:
                seen[lbl].add(red_txt)
                collected[lbl].append(red_txt)

    return {k: v for k, v in collected.items() if v}

def extract_red_text(input_doc):
    # input_doc: docx.Document object or file path
    if isinstance(input_doc, str):
        doc = Document(input_doc)
    else:
        doc = input_doc
    out = {}
    table_count = 0
    for tbl in doc.tables:
        table_count += 1
        multi_schemas = check_multi_schema_table(tbl)
        if multi_schemas:
            multi_data = extract_multi_schema_table(tbl, multi_schemas)
            for schema_name, schema_data in multi_data.items():
                if schema_data:
                    if schema_name in out:
                        for k, v in schema_data.items():
                            if k in out[schema_name]:
                                out[schema_name][k].extend(v)
                            else:
                                out[schema_name][k] = v
                    else:
                        out[schema_name] = schema_data
            continue
        schema = match_table_schema(tbl)
        if not schema:
            continue
        spec = TABLE_SCHEMAS[schema]
        data = extract_table_data(tbl, schema, spec)
        if data:
            if schema in out:
                for k, v in data.items():
                    if k in out[schema]:
                        out[schema][k].extend(v)
                    else:
                        out[schema][k] = v
            else:
                out[schema] = data

    # paragraphs (FIX: do not return early; build full 'paras' then attach)
    paras = {}
    for idx, para in enumerate(doc.paragraphs):
        red_txt = "".join(r.text for r in para.runs if is_red_font(r)).strip()
        if not red_txt:
            continue
        context = None
        for j in range(idx-1, -1, -1):
            txt = normalize_text(doc.paragraphs[j].text)
            if txt:
                all_patterns = HEADING_PATTERNS["main"] + HEADING_PATTERNS["sub"]
                if any(re.search(p, txt, re.IGNORECASE) for p in all_patterns):
                    context = txt
                    break
        if not context and re.fullmatch(PARAGRAPH_PATTERNS["date_line"], red_txt):
            context = "Date"
        if not context:
            context = "(para)"
        paras.setdefault(context, []).append(red_txt)

    if paras:
        out["paragraphs"] = paras

    # Fallback: ensure we capture the last-page Operator Declaration by headers
    if "Operator Declaration" not in out:
        op_dec = extract_operator_declaration_by_headers_from_end(doc)
        if op_dec:
            out["Operator Declaration"] = op_dec

    return out

def extract_red_text_filelike(input_file, output_file):
    """
    Accepts:
      input_file: file-like object (BytesIO/File) or path
      output_file: file-like object (opened for writing text) or path
    """
    if hasattr(input_file, "seek"):
        input_file.seek(0)
    doc = Document(input_file)
    result = extract_red_text(doc)
    if hasattr(output_file, "write"):
        json.dump(result, output_file, indent=2, ensure_ascii=False)
        output_file.flush()
    else:
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
    return result

if __name__ == "__main__":
    # Support both script and app/file-like usage
    if len(sys.argv) == 3:
        input_docx = sys.argv[1]
        output_json = sys.argv[2]
        doc = Document(input_docx)
        word_data = extract_red_text(doc)
        with open(output_json, 'w', encoding='utf-8') as f:
            json.dump(word_data, f, indent=2, ensure_ascii=False)
        print(json.dumps(word_data, indent=2, ensure_ascii=False))
    else:
        print("To use as a module: extract_red_text_filelike(input_file, output_file)")
